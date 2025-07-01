import os
import streamlit as st
from dotenv import load_dotenv
from openai import OpenAI
import google.generativeai as genai
import requests
from serpapi import GoogleSearch
from PyPDF2 import PdfReader
import docx
from langchain_community.vectorstores import FAISS
from langchain.embeddings import OpenAIEmbeddings
from langchain.text_splitter import CharacterTextSplitter
from langchain.docstore.document import Document

# Load API keys từ Streamlit Secrets hoặc biến môi trường
load_dotenv()
OPENAI_API_KEY = st.secrets.get("OPENAI_API_KEY", os.getenv("OPENAI_API_KEY"))
GEMINI_API_KEY = st.secrets.get("GEMINI_API_KEY", os.getenv("GEMINI_API_KEY"))
DEEPSEEK_API_KEY = st.secrets.get("DEEPSEEK_API_KEY", os.getenv("DEEPSEEK_API_KEY"))
SERPAPI_KEY = st.secrets.get("SERPAPI_API_KEY", os.getenv("SERPAPI_API_KEY"))

# Cấu hình API
client = OpenAI(api_key=OPENAI_API_KEY)
genai.configure(api_key=GEMINI_API_KEY)

# --- Xử lý tài liệu ---
def read_file(uploaded_file):
    if uploaded_file.type == "application/pdf":
        pdf_reader = PdfReader(uploaded_file)
        return "\n".join([page.extract_text() or "" for page in pdf_reader.pages])
    elif uploaded_file.type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
        doc = docx.Document(uploaded_file)
        return "\n".join([para.text for para in doc.paragraphs])
    elif uploaded_file.type == "text/plain":
        return str(uploaded_file.read(), "utf-8")
    else:
        return "Unsupported file format."

# --- Xử lý truy vấn SerpAPI ---
def web_search(query):
    params = {
        "engine": "google",
        "q": query,
        "api_key": SERPAPI_KEY,
        "num": 5
    }
    search = GoogleSearch(params)
    results = search.get_dict().get("organic_results", [])
    output = ""
    for idx, result in enumerate(results, 1):
        link = result.get("link", "")
        title = result.get("title", "No title")
        output += f"{idx}. [{title}]({link})\n\n"
    return output or "Không tìm thấy kết quả phù hợp."

# --- Giao diện Streamlit ---
st.set_page_config(page_title="Trợ lý AI Tổng hợp", layout="wide")
st.title("🤖 Trợ lý AI - ChatGPT x Gemini x DeepSeek x SerpAPI")

uploaded_file = st.file_uploader("Tải lên tài liệu để trợ lý AI tự học", type=["pdf", "docx", "txt"])
query = st.text_area("Nhập câu hỏi cần trợ lý AI giải quyết")

use_gpt = st.checkbox("Sử dụng ChatGPT", value=True)
use_gemini = st.checkbox("Sử dụng Gemini", value=True)
use_deepseek = st.checkbox("Sử dụng DeepSeek", value=True)
use_web = st.checkbox("Tìm kiếm trên web (SerpAPI)", value=False)
compare_mode = st.checkbox("Chế độ so sánh")

if st.button("Gửi yêu cầu"):
    docs = []
    if uploaded_file:
        content = read_file(uploaded_file)
        text_splitter = CharacterTextSplitter(chunk_size=500, chunk_overlap=50)
        chunks = text_splitter.split_text(content)
        docs = [Document(page_content=c) for c in chunks]
        st.success(f"✅ Đã đọc {len(docs)} đoạn văn bản từ tài liệu.")
        embeddings = OpenAIEmbeddings(openai_api_key=OPENAI_API_KEY)
        vectorstore = FAISS.from_documents(docs, embedding=embeddings)
        retriever = vectorstore.as_retriever()
        retrieved_docs = retriever.get_relevant_documents(query)
        context = "\n".join([doc.page_content for doc in retrieved_docs[:3]])
    else:
        context = ""

    results = {}

    if use_gpt:
        gpt_prompt = f"Trả lời câu hỏi sau dựa trên ngữ cảnh:\n\n{context}\n\nCâu hỏi: {query}"
        completion = client.chat.completions.create(
            model="gpt-3.5-turbo",
            messages=[{"role": "user", "content": gpt_prompt}]
        )
        results["ChatGPT"] = completion.choices[0].message.content

    if use_gemini:
        gemini_prompt = f"Trả lời câu hỏi sau dựa trên ngữ cảnh:\n\n{context}\n\nCâu hỏi: {query}"
        gemini_model = genai.GenerativeModel("gemini-pro")
        response = gemini_model.generate_content(gemini_prompt)
        results["Gemini"] = response.text

    if use_deepseek:
        deep_headers = {
            "Authorization": f"Bearer {DEEPSEEK_API_KEY}",
            "Content-Type": "application/json"
        }
        deep_payload = {
            "model": "deepseek-chat",
            "messages": [{"role": "user", "content": f"{context}\n\n{query}"}]
        }
        r = requests.post("https://api.deepseek.com/chat/completions", headers=deep_headers, json=deep_payload)
        if r.status_code == 200:
            results["DeepSeek"] = r.json()["choices"][0]["message"]["content"]
        else:
            results["DeepSeek"] = f"Lỗi DeepSeek: {r.text}"

    if use_web:
        results["Web Search"] = web_search(query)

    if compare_mode:
        st.subheader("🧠 So sánh kết quả các AI")
        for name, answer in results.items():
            st.markdown(f"### {name}")
            st.write(answer)
            st.markdown("---")
    else:
        for name, answer in results.items():
            st.subheader(f"📌 Trả lời từ {name}")
            st.write(answer)
            break
